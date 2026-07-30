"""
Извлечение текста из загруженных файлов разных форматов.

Поддерживаются: .txt, .pdf, .docx
Работает как с обычными путями (Path), так и с файловыми объектами
из Streamlit (st.file_uploader), у которых есть .name и .read().
"""
from __future__ import annotations

import io
from pathlib import Path
from typing import Union

from pypdf import PdfReader
from docx import Document as DocxDocument

FileLike = Union[str, Path, "io.BytesIO"]


class UnsupportedFileType(Exception):
    """Формат файла не поддерживается."""


def _get_bytes_and_name(file_obj) -> tuple[bytes, str]:
    """
    Универсально достаёт (bytes, filename) как из пути на диске,
    так и из объекта st.file_uploader (UploadedFile).
    """
    if isinstance(file_obj, (str, Path)):
        path = Path(file_obj)
        return path.read_bytes(), path.name

    # Streamlit UploadedFile или BytesIO с атрибутом .name
    name = getattr(file_obj, "name", "unknown")
    # UploadedFile нужно сначала перемотать в начало
    if hasattr(file_obj, "seek"):
        file_obj.seek(0)
    data = file_obj.read()
    return data, name


def extract_text_from_txt(data: bytes) -> str:
    for encoding in ("utf-8", "cp1251", "utf-16"):
        try:
            return data.decode(encoding)
        except UnicodeDecodeError:
            continue
    # последняя попытка — игнорировать ошибки, чтобы не падать
    return data.decode("utf-8", errors="ignore")


def extract_text_from_pdf(data: bytes) -> str:
    reader = PdfReader(io.BytesIO(data))
    pages_text = []
    for page in reader.pages:
        text = page.extract_text() or ""
        pages_text.append(text)
    return "\n".join(pages_text)


def extract_text_from_docx(data: bytes) -> str:
    doc = DocxDocument(io.BytesIO(data))
    paragraphs = [p.text for p in doc.paragraphs]

    # Не забываем про таблицы — в НПА часто важные данные лежат в них
    for table in doc.tables:
        for row in table.rows:
            row_text = " | ".join(cell.text.strip() for cell in row.cells)
            if row_text.strip(" |"):
                paragraphs.append(row_text)

    return "\n".join(paragraphs)


def extract_text(file_obj) -> tuple[str, str]:
    """
    Главная функция: принимает файл (путь или st.file_uploader объект),
    возвращает (текст, имя_файла).
    """
    data, name = _get_bytes_and_name(file_obj)
    suffix = Path(name).suffix.lower()

    if suffix == ".txt":
        text = extract_text_from_txt(data)
    elif suffix == ".pdf":
        text = extract_text_from_pdf(data)
    elif suffix == ".docx":
        text = extract_text_from_docx(data)
    else:
        raise UnsupportedFileType(
            f"Формат '{suffix}' не поддерживается. Используйте .txt, .pdf или .docx"
        )

    return text, name
